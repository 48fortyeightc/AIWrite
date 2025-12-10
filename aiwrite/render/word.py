"""
Word å¯¼å‡ºå™¨

å°†è®ºæ–‡å¯¼å‡ºä¸ºè§„èŒƒçš„ Word æ–‡æ¡£
"""

from __future__ import annotations

import subprocess
import shutil
import re
from pathlib import Path
from typing import Literal, TYPE_CHECKING

from rich.console import Console

from ..models import Paper, Section, Figure, Table
from .latex import LatexRenderer

if TYPE_CHECKING:
    from docx import Document

console = Console()


class WordExporter:
    """
    Word å¯¼å‡ºå™¨
    
    å°†è®ºæ–‡å¯¼å‡ºä¸º Word (.docx) æ–‡æ¡£
    æ”¯æŒä¸¤ç§æ–¹å¼ï¼š
    1. LaTeX â†’ Word (é€šè¿‡ pandoc)
    2. ç›´æ¥ç”Ÿæˆ Word (é€šè¿‡ python-docx)
    """

    def __init__(
        self,
        method: Literal["pandoc", "docx"] = "docx",
        pandoc_path: str | None = None,
        reference_doc: str | Path | None = None,
        images_base_path: str | Path | None = None,
    ):
        """
        åˆå§‹åŒ–å¯¼å‡ºå™¨
        
        Args:
            method: å¯¼å‡ºæ–¹æ³•ï¼Œ"pandoc" æˆ– "docx"ï¼ˆé»˜è®¤æ”¹ä¸º docxï¼‰
            pandoc_path: pandoc å¯æ‰§è¡Œæ–‡ä»¶è·¯å¾„
            reference_doc: Word å‚è€ƒæ¨¡æ¿æ–‡ä»¶ï¼ˆç”¨äºæ ·å¼ï¼‰
            images_base_path: å›¾ç‰‡æ–‡ä»¶çš„åŸºç¡€è·¯å¾„
        """
        self.method = method
        self.pandoc_path = pandoc_path or self._find_pandoc()
        self.reference_doc = Path(reference_doc) if reference_doc else None
        self.images_base_path = Path(images_base_path) if images_base_path else None
        self.latex_renderer = LatexRenderer()

    def _resolve_image_path(self, figure_path: str) -> Path:
        """
        æ™ºèƒ½è§£æå›¾ç‰‡è·¯å¾„ï¼Œé¿å…è·¯å¾„é‡å¤æ‹¼æ¥
        
        ä¾‹å¦‚ï¼š
        - images_base_path = D:/project/img
        - figure_path = img/1.png
        - ç»“æœåº”ä¸º D:/project/img/1.png è€Œä¸æ˜¯ D:/project/img/img/1.png
        """
        figure_path_obj = Path(figure_path)
        
        if not self.images_base_path:
            return figure_path_obj
        
        # ç›´æ¥æ‹¼æ¥çš„è·¯å¾„
        direct_path = self.images_base_path / figure_path
        
        # å¦‚æœç›´æ¥æ‹¼æ¥çš„è·¯å¾„å­˜åœ¨ï¼Œä½¿ç”¨å®ƒ
        if direct_path.exists():
            return direct_path
        
        # æ£€æŸ¥æ˜¯å¦æœ‰è·¯å¾„é‡å¤ï¼ˆå¦‚ base=.../img, path=img/1.pngï¼‰
        # å°è¯•åªä½¿ç”¨æ–‡ä»¶å
        filename_only = self.images_base_path / figure_path_obj.name
        if filename_only.exists():
            return filename_only
        
        # å°è¯•å»æ‰ figure_path çš„ç¬¬ä¸€å±‚ç›®å½•
        parts = figure_path_obj.parts
        if len(parts) > 1:
            without_first = Path(*parts[1:])
            alt_path = self.images_base_path / without_first
            if alt_path.exists():
                return alt_path
        
        # éƒ½ä¸å­˜åœ¨ï¼Œè¿”å›ç›´æ¥æ‹¼æ¥çš„è·¯å¾„ï¼ˆè®©åç»­ä»£ç å¤„ç†ä¸å­˜åœ¨çš„æƒ…å†µï¼‰
        return direct_path

    def _find_pandoc(self) -> str | None:
        """æŸ¥æ‰¾ pandoc å¯æ‰§è¡Œæ–‡ä»¶"""
        pandoc_path = shutil.which("pandoc")
        return pandoc_path

    def check_pandoc(self) -> bool:
        """æ£€æŸ¥ pandoc æ˜¯å¦å¯ç”¨"""
        if not self.pandoc_path:
            return False
        try:
            result = subprocess.run(
                [self.pandoc_path, "--version"],
                capture_output=True,
                text=True,
            )
            return result.returncode == 0
        except Exception:
            return False

    def export(
        self,
        paper: Paper,
        output_path: str | Path,
        use_final: bool = True,
    ) -> Path:
        """
        å¯¼å‡ºè®ºæ–‡ä¸º Word æ–‡æ¡£
        
        Args:
            paper: è®ºæ–‡å¯¹è±¡
            output_path: è¾“å‡ºæ–‡ä»¶è·¯å¾„
            use_final: æ˜¯å¦ä½¿ç”¨æ¶¦è‰²åçš„å†…å®¹
            
        Returns:
            è¾“å‡ºæ–‡ä»¶è·¯å¾„
        """
        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)

        if self.method == "pandoc":
            return self._export_via_pandoc(paper, output_path, use_final)
        else:
            return self._export_via_docx(paper, output_path, use_final)

    def _export_via_pandoc(
        self,
        paper: Paper,
        output_path: Path,
        use_final: bool,
    ) -> Path:
        """é€šè¿‡ pandoc å¯¼å‡º"""
        if not self.check_pandoc():
            raise RuntimeError(
                "pandoc æœªå®‰è£…æˆ–ä¸å¯ç”¨ã€‚è¯·å…ˆå®‰è£… pandoc: https://pandoc.org/installing.html"
            )

        console.print("[cyan]ğŸ“„ æ­£åœ¨ç”Ÿæˆ LaTeX æ–‡ä»¶...[/cyan]")

        # å…ˆç”Ÿæˆ LaTeX æ–‡ä»¶
        latex_path = output_path.with_suffix(".tex")
        self.latex_renderer.render_to_file(paper, latex_path, use_final)

        console.print("[cyan]ğŸ”„ æ­£åœ¨è½¬æ¢ä¸º Word æ–‡æ¡£...[/cyan]")

        # æ„å»º pandoc å‘½ä»¤
        cmd = [
            self.pandoc_path,
            str(latex_path),
            "-o", str(output_path),
            "--from", "latex",
            "--to", "docx",
        ]

        if self.reference_doc and self.reference_doc.exists():
            cmd.extend(["--reference-doc", str(self.reference_doc)])

        # æ‰§è¡Œè½¬æ¢
        result = subprocess.run(
            cmd,
            capture_output=True,
            text=True,
            cwd=latex_path.parent,
        )

        if result.returncode != 0:
            console.print(f"[red]pandoc é”™è¯¯: {result.stderr}[/red]")
            raise RuntimeError(f"pandoc è½¬æ¢å¤±è´¥: {result.stderr}")

        console.print(f"[green]âœ“ Word æ–‡æ¡£å·²ç”Ÿæˆ: {output_path}[/green]")

        return output_path

    def _export_via_docx(
        self,
        paper: Paper,
        output_path: Path,
        use_final: bool,
    ) -> Path:
        """é€šè¿‡ python-docx ç›´æ¥ç”Ÿæˆ Word"""
        try:
            from docx import Document
            from docx.shared import Pt, Inches, Cm
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.enum.style import WD_STYLE_TYPE
            from docx.oxml.ns import qn
            from docx.oxml import OxmlElement
        except ImportError:
            raise RuntimeError(
                "python-docx æœªå®‰è£…ã€‚è¯·è¿è¡Œ: pip install python-docx"
            )

        console.print("[cyan]ğŸ“„ æ­£åœ¨ç›´æ¥ç”Ÿæˆ Word æ–‡æ¡£...[/cyan]")

        doc = Document()
        
        # è®¾ç½®æ–‡æ¡£é»˜è®¤å­—ä½“å’Œæ®µè½æ ¼å¼
        self._set_document_defaults(doc)

        # è®ºæ–‡æ ‡é¢˜
        title_para = doc.add_paragraph()
        title_run = title_para.add_run(paper.title)
        title_run.bold = True
        title_run.font.size = Pt(22)
        title_run.font.name = 'é»‘ä½“'
        title_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'é»‘ä½“')
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_para.paragraph_format.first_line_indent = Cm(0)  # æ ‡é¢˜ä¸ç¼©è¿›
        title_para.space_after = Pt(24)

        # ä½œè€…
        if paper.authors:
            author_para = doc.add_paragraph()
            author_run = author_para.add_run(", ".join(paper.authors))
            author_run.font.size = Pt(12)
            author_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            author_para.paragraph_format.first_line_indent = Cm(0)

        doc.add_page_break()

        # æ·»åŠ ç›®å½•æ ‡é¢˜
        toc_title = doc.add_paragraph()
        toc_run = toc_title.add_run("ç›®  å½•")
        toc_run.bold = True
        toc_run.font.size = Pt(16)
        toc_run.font.name = 'é»‘ä½“'
        toc_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'é»‘ä½“')
        toc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        toc_title.paragraph_format.first_line_indent = Cm(0)
        toc_title.space_after = Pt(18)

        # æ·»åŠ ç›®å½•åŸŸï¼ˆéœ€è¦ç”¨æˆ·æ‰‹åŠ¨æ›´æ–°ï¼‰
        self._add_toc_field(doc)
        
        doc.add_paragraph()  # ç©ºè¡Œ
        hint_para = doc.add_paragraph()
        hint_run = hint_para.add_run('ï¼ˆè¯·å³é”®ç‚¹å‡»ç›®å½•ï¼Œé€‰æ‹©"æ›´æ–°åŸŸ"ä»¥ç”Ÿæˆç›®å½•ï¼‰')
        hint_run.italic = True
        hint_para.paragraph_format.first_line_indent = Cm(0)

        # ç« èŠ‚å†…å®¹
        for i, section in enumerate(paper.sections):
            # æ¯ä¸ªä¸»è¦ç« èŠ‚å‰æ·»åŠ åˆ†é¡µç¬¦
            doc.add_page_break()
            self._add_section_to_doc(doc, section, use_final, level=1, is_first_section=(i==0), keywords=paper.keywords, keywords_en=paper.keywords_en)

        # ä¿å­˜æ–‡æ¡£
        doc.save(str(output_path))
        console.print(f"[green]âœ“ Word æ–‡æ¡£å·²ç”Ÿæˆ: {output_path}[/green]")

        return output_path

    def _set_document_defaults(self, doc: "Document") -> None:
        """è®¾ç½®æ–‡æ¡£é»˜è®¤æ ·å¼"""
        from docx.shared import Pt, Cm
        from docx.oxml.ns import qn

        # è®¾ç½®æ­£æ–‡æ ·å¼
        style = doc.styles['Normal']
        style.font.name = 'å®‹ä½“'
        style.font.size = Pt(12)
        style._element.rPr.rFonts.set(qn('w:eastAsia'), 'å®‹ä½“')
        style.paragraph_format.line_spacing = 1.5
        style.paragraph_format.first_line_indent = Cm(0.74)  # é¦–è¡Œç¼©è¿›2å­—ç¬¦

        # è®¾ç½®æ ‡é¢˜æ ·å¼
        for i in range(1, 4):
            heading_style = doc.styles[f'Heading {i}']
            if i == 1:
                heading_style.font.size = Pt(16)
            elif i == 2:
                heading_style.font.size = Pt(14)
            else:
                heading_style.font.size = Pt(12)
            heading_style.font.name = 'é»‘ä½“'
            heading_style.font.bold = True
            heading_style._element.rPr.rFonts.set(qn('w:eastAsia'), 'é»‘ä½“')
            heading_style.paragraph_format.first_line_indent = Cm(0)  # æ ‡é¢˜ä¸ç¼©è¿›
            heading_style.paragraph_format.space_before = Pt(12)
            heading_style.paragraph_format.space_after = Pt(6)

    def _add_toc_field(self, doc: "Document") -> None:
        """æ·»åŠ ç›®å½•åŸŸ"""
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn

        paragraph = doc.add_paragraph()
        run = paragraph.add_run()
        
        # åˆ›å»ºå¤æ‚åŸŸ
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
        
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'separate')
        
        fldChar3 = OxmlElement('w:fldChar')
        fldChar3.set(qn('w:fldCharType'), 'end')
        
        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)
        run._r.append(fldChar3)

    def _add_section_to_doc(
        self,
        doc: "Document",
        section: "Section",
        use_final: bool,
        level: int,
        is_first_section: bool = False,
        keywords: list[str] | None = None,
        keywords_en: list[str] | None = None,
    ) -> None:
        """æ·»åŠ ç« èŠ‚åˆ° Word æ–‡æ¡£
        
        Args:
            keywords: ä¸­æ–‡å…³é”®è¯åˆ—è¡¨ï¼Œä»…åœ¨æ‘˜è¦ç« èŠ‚åè¾“å‡º
            keywords_en: è‹±æ–‡å…³é”®è¯åˆ—è¡¨ï¼Œä»…åœ¨è‹±æ–‡æ‘˜è¦ç« èŠ‚åè¾“å‡º
        """
        from docx.shared import Pt, Cm, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn

        # è·å–å†…å®¹
        content = ""
        if use_final and section.final_latex:
            content = section.final_latex
        elif section.draft_latex:
            content = section.draft_latex

        # å§‹ç»ˆæ·»åŠ ç« èŠ‚æ ‡é¢˜ï¼ˆä¸ç®¡å†…å®¹ä¸­æ˜¯å¦æœ‰å­æ ‡é¢˜ï¼‰
        heading_level = min(level, 9)
        heading = doc.add_heading(section.title, level=heading_level)
        
        # è®¾ç½®æ ‡é¢˜å­—ä½“
        for run in heading.runs:
            run.font.name = 'é»‘ä½“'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'é»‘ä½“')
        
        heading.paragraph_format.first_line_indent = Cm(0)

        if content:
            # æ”¶é›†æœ¬ç« èŠ‚åŠæ‰€æœ‰å­ç« èŠ‚çš„å›¾ç‰‡å’Œè¡¨æ ¼
            all_figures = self._collect_all_figures(section)
            all_tables = self._collect_all_tables(section)
            # è§£æ LaTeX å†…å®¹å¹¶æ·»åŠ åˆ°æ–‡æ¡£ï¼ˆåŒ…æ‹¬å¤„ç† \subsectionï¼‰
            self._add_latex_content_to_doc(doc, content, level, all_figures, all_tables)
        
        # å¦‚æœæ˜¯æ‘˜è¦ç« èŠ‚ï¼Œåœ¨å†…å®¹åæ·»åŠ å…³é”®è¯
        if keywords and section.id in ("abstract-zh", "abstract", "æ‘˜è¦"):
            kw_para = doc.add_paragraph()
            kw_bold = kw_para.add_run("å…³é”®è¯ï¼š")
            kw_bold.bold = True
            kw_bold.font.name = 'å®‹ä½“'
            kw_bold._element.rPr.rFonts.set(qn('w:eastAsia'), 'å®‹ä½“')
            kw_run = kw_para.add_run("ï¼›".join(keywords))
            kw_run.font.name = 'å®‹ä½“'
            kw_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'å®‹ä½“')
            kw_para.paragraph_format.first_line_indent = Cm(0)
            kw_para.space_after = Pt(12)
        
        # å¦‚æœæ˜¯è‹±æ–‡æ‘˜è¦ç« èŠ‚ï¼Œåœ¨å†…å®¹åæ·»åŠ  Keywords
        if keywords and section.id in ("abstract-en", "Abstract"):
            kw_para = doc.add_paragraph()
            kw_bold = kw_para.add_run("Keywords: ")
            kw_bold.bold = True
            # ä½¿ç”¨è‹±æ–‡å…³é”®è¯ï¼ˆå¦‚æœæœ‰ï¼‰ï¼Œå¦åˆ™ä½¿ç”¨ä¸­æ–‡å…³é”®è¯
            kw_en = keywords_en if keywords_en else keywords
            kw_para.add_run("; ".join(kw_en))
            kw_para.paragraph_format.first_line_indent = Cm(0)
            kw_para.space_after = Pt(12)
        
        # å¦‚æœæ²¡æœ‰å†…å®¹ä½†æœ‰å›¾ç‰‡ï¼Œå•ç‹¬æ·»åŠ å›¾ç‰‡
        if not content and section.figures:
            self._add_figures_to_doc(doc, section.figures)

        # é€’å½’å¤„ç†å­ç« èŠ‚ï¼ˆå¦‚æœæ²¡æœ‰å†…å®¹æˆ–å†…å®¹ä¸­æ²¡æœ‰\subsectionï¼‰
        # å¦‚æœå†…å®¹ä¸­å·²ç»åŒ…å«\subsectionï¼Œè¯´æ˜å­æ ‡é¢˜å·²ç»åœ¨å†…å®¹ä¸­äº†ï¼Œä¸éœ€è¦é€’å½’
        has_subsections_in_content = bool(re.search(r"\\subsection\{", content))
        if not has_subsections_in_content:
            for child in section.children:
                self._add_section_to_doc(doc, child, use_final, level + 1, keywords=keywords, keywords_en=keywords_en)

    def _add_figures_to_doc(
        self,
        doc: "Document",
        figures: list[Figure],
    ) -> None:
        """æ·»åŠ å›¾ç‰‡åˆ°æ–‡æ¡£"""
        from docx.shared import Pt, Cm, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn

        for figure in figures:
            # ç¡®å®šå›¾ç‰‡è·¯å¾„ï¼ˆæ™ºèƒ½å¤„ç†è·¯å¾„é‡å¤ï¼‰
            image_path = self._resolve_image_path(figure.path)
            
            # æ·»åŠ å›¾ç‰‡
            if image_path.exists():
                try:
                    # æ·»åŠ å›¾ç‰‡ï¼ˆå®½åº¦ä¸ºé¡µé¢å®½åº¦çš„ 80%ï¼‰
                    para = doc.add_paragraph()
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = para.add_run()
                    run.add_picture(str(image_path), width=Inches(5))
                    
                    # æ·»åŠ å›¾ç‰‡æ ‡é¢˜
                    caption_para = doc.add_paragraph()
                    caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    caption_run = caption_para.add_run(f"å›¾ {figure.id}: {figure.caption}")
                    caption_run.font.name = 'å®‹ä½“'
                    caption_run.font.size = Pt(10)
                    caption_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'å®‹ä½“')
                    caption_para.paragraph_format.first_line_indent = Cm(0)
                    caption_para.space_after = Pt(12)
                    
                    console.print(f"[green]  âœ“ æ’å…¥å›¾ç‰‡: {figure.caption}[/green]")
                except Exception as e:
                    console.print(f"[yellow]  âš  å›¾ç‰‡æ’å…¥å¤±è´¥ {figure.path}: {e}[/yellow]")
                    # æ·»åŠ å ä½ç¬¦
                    self._add_figure_placeholder(doc, figure)
            else:
                console.print(f"[yellow]  âš  å›¾ç‰‡ä¸å­˜åœ¨: {image_path}[/yellow]")
                # æ·»åŠ å ä½ç¬¦
                self._add_figure_placeholder(doc, figure)

    def _add_figure_placeholder(
        self,
        doc: "Document",
        figure: Figure,
    ) -> None:
        """æ·»åŠ å›¾ç‰‡å ä½ç¬¦"""
        from docx.shared import Pt, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn

        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(f"[å›¾ {figure.id}: {figure.caption}]")
        run.font.name = 'å®‹ä½“'
        run.font.size = Pt(10)
        run.italic = True
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'å®‹ä½“')
        para.paragraph_format.first_line_indent = Cm(0)
        
        if figure.description:
            desc_para = doc.add_paragraph()
            desc_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            desc_run = desc_para.add_run(f"ï¼ˆ{figure.description[:100]}...ï¼‰" if len(figure.description) > 100 else f"ï¼ˆ{figure.description}ï¼‰")
            desc_run.font.name = 'å®‹ä½“'
            desc_run.font.size = Pt(9)
            desc_run.italic = True
            desc_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'å®‹ä½“')
            desc_para.paragraph_format.first_line_indent = Cm(0)

    def _insert_table(
        self,
        doc: "Document",
        table: Table,
    ) -> None:
        """æ’å…¥è¡¨æ ¼åˆ°æ–‡æ¡£"""
        from docx.shared import Pt, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.oxml.ns import qn
        
        # å¦‚æœæœ‰ pathï¼Œä» Excel è¯»å–è¡¨æ ¼å†…å®¹
        if table.path:
            try:
                from ..utils.excel import read_excel_file
                
                # ç¡®å®šè¡¨æ ¼è·¯å¾„ï¼ˆæ™ºèƒ½å¤„ç†è·¯å¾„é‡å¤ï¼‰
                table_path = self._resolve_image_path(table.path)
                
                if table_path.exists():
                    rows = read_excel_file(table_path)
                    if rows:
                        self._create_word_table(doc, rows, table)
                        console.print(f"[green]  âœ“ æ’å…¥è¡¨æ ¼: {table.caption}[/green]")
                        return
                    else:
                        console.print(f"[yellow]  âš  è¡¨æ ¼æ–‡ä»¶ä¸ºç©º: {table_path}[/yellow]")
                else:
                    console.print(f"[yellow]  âš  è¡¨æ ¼æ–‡ä»¶ä¸å­˜åœ¨: {table_path}[/yellow]")
            except Exception as e:
                console.print(f"[red]  âœ— è¯»å–è¡¨æ ¼å¤±è´¥ {table.path}: {e}[/red]")
        
        # å¦‚æœæœ‰ contentï¼ˆMarkdown æ ¼å¼ï¼‰ï¼Œè§£æå¹¶åˆ›å»ºè¡¨æ ¼
        if table.content:
            rows = self._parse_markdown_table(table.content)
            if rows:
                self._create_word_table(doc, rows, table)
                console.print(f"[green]  âœ“ æ’å…¥è¡¨æ ¼: {table.caption}[/green]")
                return
        
        # å¦åˆ™æ·»åŠ å ä½ç¬¦
        self._add_table_placeholder(doc, table)

    def _create_word_table(
        self,
        doc: "Document",
        rows: list[list[str]],
        table: Table,
    ) -> None:
        """åˆ›å»º Word è¡¨æ ¼"""
        from docx.shared import Pt, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.oxml.ns import qn
        
        if not rows:
            return
        
        num_rows = len(rows)
        num_cols = max(len(row) for row in rows)
        
        # åˆ›å»ºè¡¨æ ¼
        word_table = doc.add_table(rows=num_rows, cols=num_cols)
        word_table.style = 'Table Grid'
        word_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # å¡«å……è¡¨æ ¼å†…å®¹
        for i, row in enumerate(rows):
            for j, cell_text in enumerate(row):
                if j < num_cols:
                    cell = word_table.rows[i].cells[j]
                    cell.text = cell_text
                    
                    # è®¾ç½®å•å…ƒæ ¼å­—ä½“
                    for paragraph in cell.paragraphs:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        for run in paragraph.runs:
                            run.font.name = 'å®‹ä½“'
                            run.font.size = Pt(10)
                            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'å®‹ä½“')
                            # è¡¨å¤´åŠ ç²—
                            if i == 0:
                                run.bold = True
        
        # æ·»åŠ è¡¨æ ¼æ ‡é¢˜ï¼ˆåœ¨è¡¨æ ¼ä¸‹æ–¹ï¼‰
        caption_para = doc.add_paragraph()
        caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption_run = caption_para.add_run(f"è¡¨ {table.id}: {table.caption}")
        caption_run.font.name = 'å®‹ä½“'
        caption_run.font.size = Pt(10)
        caption_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'å®‹ä½“')
        caption_para.paragraph_format.first_line_indent = Cm(0)
        caption_para.space_after = Pt(12)

    def _parse_markdown_table(self, content: str) -> list[list[str]]:
        """è§£æ Markdown æ ¼å¼çš„è¡¨æ ¼"""
        lines = content.strip().split('\n')
        rows: list[list[str]] = []
        
        for line in lines:
            line = line.strip()
            if not line or line.startswith('|--') or line.startswith('| --'):
                continue  # è·³è¿‡åˆ†éš”è¡Œ
            if line.startswith('|'):
                # è§£æè¡¨æ ¼è¡Œ
                cells = [cell.strip() for cell in line.split('|')[1:-1]]
                if cells:
                    rows.append(cells)
        
        return rows

    def _add_table_placeholder(
        self,
        doc: "Document",
        table: Table,
    ) -> None:
        """æ·»åŠ è¡¨æ ¼å ä½ç¬¦"""
        from docx.shared import Pt, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn

        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(f"[è¡¨ {table.id}: {table.caption}]")
        run.font.name = 'å®‹ä½“'
        run.font.size = Pt(10)
        run.italic = True
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'å®‹ä½“')
        para.paragraph_format.first_line_indent = Cm(0)
        
        if table.description:
            desc_para = doc.add_paragraph()
            desc_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            desc_run = desc_para.add_run(f"ï¼ˆ{table.description[:100]}...ï¼‰" if len(table.description) > 100 else f"ï¼ˆ{table.description}ï¼‰")
            desc_run.font.name = 'å®‹ä½“'
            desc_run.font.size = Pt(9)
            desc_run.italic = True
            desc_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'å®‹ä½“')
            desc_para.paragraph_format.first_line_indent = Cm(0)

    def _collect_all_figures(self, section: Section) -> list[Figure]:
        """é€’å½’æ”¶é›†ç« èŠ‚åŠå…¶æ‰€æœ‰å­ç« èŠ‚çš„å›¾ç‰‡"""
        all_figures: list[Figure] = []
        
        # æ·»åŠ å½“å‰ç« èŠ‚çš„å›¾ç‰‡
        all_figures.extend(section.figures)
        
        # é€’å½’æ·»åŠ å­ç« èŠ‚çš„å›¾ç‰‡
        for child in section.children:
            all_figures.extend(self._collect_all_figures(child))
        
        return all_figures

    def _collect_all_tables(self, section: Section) -> list[Table]:
        """é€’å½’æ”¶é›†ç« èŠ‚åŠå…¶æ‰€æœ‰å­ç« èŠ‚çš„è¡¨æ ¼"""
        all_tables: list[Table] = []
        
        # æ·»åŠ å½“å‰ç« èŠ‚çš„è¡¨æ ¼
        all_tables.extend(section.tables)
        
        # é€’å½’æ·»åŠ å­ç« èŠ‚çš„è¡¨æ ¼
        for child in section.children:
            all_tables.extend(self._collect_all_tables(child))
        
        return all_tables

    def _add_latex_content_to_doc(
        self,
        doc: "Document",
        latex_content: str,
        base_level: int = 1,
        figures: list[Figure] | None = None,
        tables: list[Table] | None = None,
    ) -> None:
        """å°† LaTeX å†…å®¹è½¬æ¢å¹¶æ·»åŠ åˆ° Word æ–‡æ¡£"""
        from docx.shared import Pt, Cm, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn

        figures = figures or []
        tables = tables or []
        figure_map = {f.id: f for f in figures}
        # åŒæ—¶ç”¨ caption ä½œä¸ºå¤‡ç”¨ key
        for f in figures:
            figure_map[f.caption] = f
        table_map = {t.id: t for t in tables}
        for t in tables:
            table_map[t.caption] = t
        figures_inserted = set()
        tables_inserted = set()

        # å…ˆå¤„ç† LaTeX è½¬ä¹‰ç¬¦å·
        content = latex_content
        content = content.replace("\\%", "%")
        content = content.replace("\\$", "$")
        content = content.replace("\\&", "&")
        content = content.replace("\\#", "#")
        content = content.replace("\\_", "_")
        content = content.replace("\\{", "{")
        content = content.replace("\\}", "}")
        
        # å¤„ç†æ•°å­¦å…¬å¼ $...$
        content = re.sub(r"\$([^$]+)\$", r"\1", content)
        
        # å¤„ç†å›¾è¡¨å ä½ç¬¦ - è½¬æ¢ä¸ºç‰¹æ®Šæ ‡è®°ä»¥ä¾¿åç»­å¤„ç†
        content = re.sub(r"\{\{FIGURE:([^:]*):([^}]*)\}\}", r"\n\n<<FIGURE:\1:\2>>\n\n", content)
        content = re.sub(r"\{\{TABLE:([^:]*):([^}]*)\}\}", r"\n\n<<TABLE:\1:\2>>\n\n", content)
        
        # å¤„ç† \textbf{...} - ä¿ç•™å†…å®¹
        content = re.sub(r"\\textbf\{([^}]*)\}", r"\1", content)
        
        # å¤„ç† \textit{...} å’Œ \emph{...}
        content = re.sub(r"\\textit\{([^}]*)\}", r"\1", content)
        content = re.sub(r"\\emph\{([^}]*)\}", r"\1", content)
        
        # å¤„ç† \cite{...}
        content = re.sub(r"\\cite\{[^}]*\}", "[å¼•ç”¨]", content)
        
        # ç§»é™¤ \label{...} å’Œ \ref{...} ç­‰å¼•ç”¨å‘½ä»¤ï¼ˆè¿™äº›ä¸åº”è¯¥å‡ºç°åœ¨Wordä¸­ï¼‰
        content = re.sub(r"\\label\{[^}]*\}", "", content)
        content = re.sub(r"\\ref\{[^}]*\}", "", content)
        content = re.sub(r"\\pageref\{[^}]*\}", "", content)
        content = re.sub(r"\\autoref\{[^}]*\}", "", content)
        
        # ä½¿ç”¨ç‰¹æ®Šæ ‡è®°åˆ†å‰²ç« èŠ‚å’Œå†…å®¹
        # å°† \section{...} \subsection{...} ç­‰æ›¿æ¢ä¸ºç‰¹æ®Šæ ‡è®° (LaTeXæ ¼å¼)
        content = re.sub(r"\\section\{([^}]*)\}", r"\n\n<<HEADING:1:\1>>\n\n", content)
        content = re.sub(r"\\subsection\{([^}]*)\}", r"\n\n<<HEADING:2:\1>>\n\n", content)
        content = re.sub(r"\\subsubsection\{([^}]*)\}", r"\n\n<<HEADING:3:\1>>\n\n", content)
        
        # åŒæ—¶å¤„ç† Markdown æ ¼å¼çš„æ ‡é¢˜
        # ### ä¸‰çº§æ ‡é¢˜
        content = re.sub(r"^### (.+)$", r"\n\n<<HEADING:3:\1>>\n\n", content, flags=re.MULTILINE)
        # ## äºŒçº§æ ‡é¢˜
        content = re.sub(r"^## (.+)$", r"\n\n<<HEADING:2:\1>>\n\n", content, flags=re.MULTILINE)
        # # ä¸€çº§æ ‡é¢˜ (ç« èŠ‚ä¸»æ ‡é¢˜,é€šå¸¸å·²ç»åœ¨å¤–é¢å¤„ç†äº†,è¿™é‡Œå¿½ç•¥)
        content = re.sub(r"^# .+$", r"", content, flags=re.MULTILINE)
        
        # ç§»é™¤å…¶ä»– LaTeX å‘½ä»¤ä½†ä¿ç•™å†…å®¹
        content = re.sub(r"\\[a-zA-Z]+\*?\{([^}]*)\}", r"\1", content)
        content = re.sub(r"\\[a-zA-Z]+\*?", "", content)
        content = re.sub(r"[{}]", "", content)
        
        # ç§»é™¤æ®‹ç•™çš„ sec: subsec: ç­‰æ ‡ç­¾æ–‡æœ¬
        content = re.sub(r"\b(sec|subsec|fig|tab|eq|chap):[a-zA-Z0-9_-]+\b", "", content)
        
        # æ¸…ç†å¤šä½™ç©ºè¡Œ
        content = re.sub(r"\n{3,}", "\n\n", content)
        
        # æŒ‰æ®µè½åˆ†å‰²å¹¶æ·»åŠ 
        paragraphs = content.strip().split("\n\n")
        for para_text in paragraphs:
            para_text = para_text.strip()
            
            # è·³è¿‡ç©ºæ®µè½
            if not para_text or len(para_text) < 2:
                continue
            
            # æ£€æŸ¥æ˜¯å¦æ˜¯æ ‡é¢˜æ ‡è®°
            heading_match = re.match(r"<<HEADING:(\d+):(.+)>>", para_text)
            if heading_match:
                heading_rel_level = int(heading_match.group(1))
                heading_title = heading_match.group(2).strip()
                # è®¡ç®—å®é™…çš„æ ‡é¢˜çº§åˆ«
                actual_level = base_level + heading_rel_level - 1
                actual_level = min(actual_level, 9)
                
                heading = doc.add_heading(heading_title, level=actual_level)
                for run in heading.runs:
                    run.font.name = 'é»‘ä½“'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'é»‘ä½“')
                heading.paragraph_format.first_line_indent = Cm(0)
                continue
            
            # æ£€æŸ¥æ˜¯å¦æ˜¯å›¾ç‰‡æ ‡è®°
            figure_match = re.match(r"<<FIGURE:([^:]*):([^>]*)>>", para_text)
            if figure_match:
                fig_caption = figure_match.group(1).strip()
                fig_desc = figure_match.group(2).strip()
                
                # å°è¯•æŸ¥æ‰¾åŒ¹é…çš„çœŸå®å›¾ç‰‡
                matched_figure = None
                for fig_id, fig in figure_map.items():
                    if fig.caption == fig_caption or fig_id not in figures_inserted:
                        matched_figure = fig
                        figures_inserted.add(fig_id)
                        break
                
                if matched_figure:
                    # æ’å…¥çœŸå®å›¾ç‰‡
                    self._insert_figure(doc, matched_figure)
                else:
                    # æ·»åŠ å ä½ç¬¦
                    para = doc.add_paragraph()
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = para.add_run(f"[å›¾: {fig_caption}]")
                    run.font.name = 'å®‹ä½“'
                    run.font.size = Pt(10)
                    run.italic = True
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'å®‹ä½“')
                    para.paragraph_format.first_line_indent = Cm(0)
                    
                    if fig_desc:
                        desc_para = doc.add_paragraph()
                        desc_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        desc_run = desc_para.add_run(f"è¯´æ˜: {fig_desc}")
                        desc_run.font.name = 'å®‹ä½“'
                        desc_run.font.size = Pt(9)
                        desc_run.italic = True
                        desc_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'å®‹ä½“')
                        desc_para.paragraph_format.first_line_indent = Cm(0)
                continue
            
            # æ£€æŸ¥æ˜¯å¦æ˜¯è¡¨æ ¼æ ‡è®°
            table_match = re.match(r"<<TABLE:([^:]*):([^>]*)>>", para_text)
            if table_match:
                tab_caption = table_match.group(1).strip()
                tab_desc = table_match.group(2).strip()
                
                # å°è¯•æŸ¥æ‰¾åŒ¹é…çš„çœŸå®è¡¨æ ¼
                matched_table = None
                for tab_key, tab in table_map.items():
                    if tab.caption == tab_caption or tab_key not in tables_inserted:
                        matched_table = tab
                        tables_inserted.add(tab.id)
                        tables_inserted.add(tab.caption)
                        break
                
                if matched_table:
                    # æ’å…¥çœŸå®è¡¨æ ¼
                    self._insert_table(doc, matched_table)
                else:
                    # æ·»åŠ å ä½ç¬¦
                    para = doc.add_paragraph()
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = para.add_run(f"[è¡¨: {tab_caption}]")
                    run.font.name = 'å®‹ä½“'
                    run.font.size = Pt(10)
                    run.italic = True
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'å®‹ä½“')
                    para.paragraph_format.first_line_indent = Cm(0)
                    
                    if tab_desc:
                        desc_para = doc.add_paragraph()
                        desc_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        desc_run = desc_para.add_run(f"è¯´æ˜: {tab_desc}")
                        desc_run.font.name = 'å®‹ä½“'
                        desc_run.font.size = Pt(9)
                        desc_run.italic = True
                        desc_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'å®‹ä½“')
                        desc_para.paragraph_format.first_line_indent = Cm(0)
                continue
            
            # åˆå¹¶æ®µè½å†…çš„æ¢è¡Œ
            para_text = re.sub(r"\s*\n\s*", " ", para_text)
            para_text = re.sub(r"\s+", " ", para_text).strip()
            
            # æ™®é€šæ®µè½
            para = doc.add_paragraph()
            run = para.add_run(para_text)
            run.font.name = 'å®‹ä½“'
            run.font.size = Pt(12)
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'å®‹ä½“')
            para.paragraph_format.first_line_indent = Cm(0.74)  # é¦–è¡Œç¼©è¿›
            para.paragraph_format.line_spacing = 1.5
        
        # æ’å…¥å‰©ä½™æœªæ’å…¥çš„å›¾ç‰‡ï¼ˆåœ¨å†…å®¹æœ«å°¾ï¼‰
        for fig in figures:
            if fig.id not in figures_inserted and fig.caption not in figures_inserted:
                self._insert_figure(doc, fig)
                figures_inserted.add(fig.id)
                figures_inserted.add(fig.caption)
        
        # æ’å…¥å‰©ä½™æœªæ’å…¥çš„è¡¨æ ¼ï¼ˆåœ¨å†…å®¹æœ«å°¾ï¼‰
        for tab in tables:
            if tab.id not in tables_inserted and tab.caption not in tables_inserted:
                self._insert_table(doc, tab)
                tables_inserted.add(tab.id)
                tables_inserted.add(tab.caption)

    def _insert_figure(
        self,
        doc: "Document",
        figure: Figure,
    ) -> None:
        """æ’å…¥å•å¼ å›¾ç‰‡"""
        from docx.shared import Pt, Cm, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn

        # æ£€æŸ¥ path æ˜¯å¦å­˜åœ¨ä¸”æœ‰æ•ˆ
        if not figure.path or figure.path.strip() in ('', '.', '..'):
            console.print(f"[yellow]  âš  å›¾ç‰‡æœªæŒ‡å®šè·¯å¾„: {figure.caption}[/yellow]")
            self._add_figure_placeholder(doc, figure)
            return

        # ç¡®å®šå›¾ç‰‡è·¯å¾„ï¼ˆæ™ºèƒ½å¤„ç†è·¯å¾„é‡å¤ï¼‰
        image_path = self._resolve_image_path(figure.path)
        
        if image_path.exists():
            try:
                # æ·»åŠ å›¾ç‰‡
                para = doc.add_paragraph()
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = para.add_run()
                run.add_picture(str(image_path), width=Inches(5))
                
                # æ·»åŠ å›¾ç‰‡æ ‡é¢˜
                caption_para = doc.add_paragraph()
                caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                caption_run = caption_para.add_run(f"å›¾ {figure.id}: {figure.caption}")
                caption_run.font.name = 'å®‹ä½“'
                caption_run.font.size = Pt(10)
                caption_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'å®‹ä½“')
                caption_para.paragraph_format.first_line_indent = Cm(0)
                caption_para.space_after = Pt(12)
                
                console.print(f"[green]  âœ“ æ’å…¥å›¾ç‰‡: {figure.caption}[/green]")
            except Exception as e:
                console.print(f"[yellow]  âš  å›¾ç‰‡æ’å…¥å¤±è´¥ {figure.path}: {e}[/yellow]")
                self._add_figure_placeholder(doc, figure)
        else:
            console.print(f"[yellow]  âš  å›¾ç‰‡ä¸å­˜åœ¨: {image_path}[/yellow]")
            self._add_figure_placeholder(doc, figure)

    def _strip_latex_commands(self, text: str) -> str:
        """ç§»é™¤ LaTeX å‘½ä»¤ï¼Œä¿ç•™æ–‡æœ¬å†…å®¹"""
        import re

        # ç§»é™¤ \section{}, \subsection{} ç­‰å‘½ä»¤
        text = re.sub(r"\\(sub)*section\{[^}]*\}", "", text)
        
        # ç§»é™¤ \textbf{...} ä½†ä¿ç•™å†…å®¹
        text = re.sub(r"\\textbf\{([^}]*)\}", r"\1", text)
        text = re.sub(r"\\textit\{([^}]*)\}", r"\1", text)
        text = re.sub(r"\\emph\{([^}]*)\}", r"\1", text)
        
        # ç§»é™¤ \cite{...}
        text = re.sub(r"\\cite\{[^}]*\}", "[å¼•ç”¨]", text)
        
        # ç§»é™¤å…¶ä»–å¸¸è§å‘½ä»¤
        text = re.sub(r"\\[a-zA-Z]+\{[^}]*\}", "", text)
        text = re.sub(r"\\[a-zA-Z]+", "", text)
        
        # æ¸…ç†å¤šä½™ç©ºç™½
        text = re.sub(r"\n{3,}", "\n\n", text)

        return text.strip()
